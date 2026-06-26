"""
build_naane_docx.py
====================

Generador de documentos Word (.docx) con el membrete oficial del proyecto
CC217 — Naãne (FUCAI / OIKOS / AICS).

REGLAS INMUTABLES (verificadas automáticamente al final):
  HEADER  : exactamente 2 logos — AICS (col 1, izquierda) + NAANE (col 2, derecha)
            ambos de TAMAÑO IDÉNTICO (2.0 cm × 2.0 cm).
  FOOTER  : exactamente 9 logos en ORDEN FIJO — OIKOS, UniTrento DICAM, FUCAI,
            Javeriana, Alcaldía Leticia, ACITAM, ATICOYA, Maikuchiga, Raíz.
            Todos de TAMAÑO IDÉNTICO (1.2 cm × 1.2 cm).

Uniformidad de tamaño:
  Todos los logos están preprocesados a canvas cuadrados (1024×1024 px) con
  el contenido centrado y fondo transparente. Al insertarlos en celdas
  cuadradas del mismo tamaño en el documento, todos aparecen del mismo
  tamaño visual sin importar la forma de su contenido original.

Validación:
  Al final del proceso, el documento se vuelve a abrir y se verifica que:
    - El header tenga exactamente 2 imágenes (AICS y NAANE).
    - El footer tenga exactamente 9 imágenes en el orden correcto.
    - Todas las imágenes del footer tengan dimensiones idénticas.
    - El bloque de título contenga el identificador 'CC217'.
  Si alguna verificación falla, se lanza NaaneBrandingValidationError.

Uso:
    from build_naane_docx import create_naane_document, validate_naane_document
    doc = create_naane_document(title="Mi informe", document_type="informe",
                                skill_dir="/path/to/naane-branding")
    # ... agregar contenido ...
    doc.save("/path/to/output.docx")
    validate_naane_document("/path/to/output.docx")  # lanza si algo falla
"""

from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===========================================================================
# CONSTANTES INMUTABLES — orden, ubicación y tamaños congelados
# ===========================================================================
# NO modificar estas constantes sin actualizar los tests de validación.

# --- Colores y tipografía FUCAI (heredados de fucai-branding) ---
NARANJA_FUCAI = RGBColor(0xE9, 0x45, 0x13)   # #E94513
ARENA         = RGBColor(0xED, 0xE8, 0xD3)   # #EDE8D3
ARENA_CLARO   = RGBColor(0xF6, 0xF3, 0xE9)   # #F6F3E9
NEGRO         = RGBColor(0x00, 0x00, 0x00)
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO    = RGBColor(0x33, 0x33, 0x33)   # #333333
GRIS_MEDIO    = RGBColor(0x66, 0x66, 0x66)   # #666666

FONT_HEADING = "Space Grotesk"
FONT_BODY    = "Calibri"

# --- Orden CANÓNICO de logos del HEADER (izquierda → derecha) ---
HEADER_LOGOS_ORDER = (
    "aics.png",      # columna 1 (izquierda) — donante: AICS
    "naane.png",     # columna 2 (derecha)   — proyecto: Naane
)

# --- Orden CANÓNICO de logos del FOOTER (izquierda → derecha) ---
FOOTER_LOGOS_ORDER = (
    "oikos.png",            # 1. OIKOS — socio italiano implementador
    "unitn_dicam.png",      # 2. Università di Trento — DICAM (académico italiano)
    "fucai.png",            # 3. FUCAI — implementador colombiano
    "javeriana.png",        # 4. Pontificia Universidad Javeriana (académico colombiano)
    "alcaldia_leticia.png", # 5. Alcaldía de Leticia (autoridad territorial)
    "acitam.png",           # 6. ACITAM (autoridad indígena)
    "aticoya.png",          # 7. ATICOYA (autoridad indígena)
    "maikuchiga.png",       # 8. Fundación Maikuchiga
    "raiz.png",             # 9. Fundación Raíz
)

# --- Tamaños FIJOS de los logos (cuadrados, idénticos) ---
HEADER_LOGO_SIDE = Cm(2.0)   # cada logo del header: 2.0 × 2.0 cm
FOOTER_LOGO_SIDE = Cm(1.2)   # cada logo del footer: 1.2 × 1.2 cm

# --- Caption fija del footer ---
FOOTER_CAPTION = (
    "Proyecto «Naãne — Buen Vivir Amazonas» (CC217)  ·  "
    "Financiado por la Agenzia Italiana per la Cooperazione allo Sviluppo (AICS)  ·  "
    "Implementado por OIKOS y Fundación Caminos de Identidad — FUCAI"
)

# --- Identificador obligatorio en todo documento NAANE ---
PROJECT_IDENTIFIER = "CC217"
PROJECT_NAME_FULL  = "Proyecto CC217 — Naãne · Buen Vivir Amazonas"


# ===========================================================================
# EXCEPCIONES
# ===========================================================================

class NaaneBrandingValidationError(RuntimeError):
    """Se lanza cuando el documento generado no cumple las reglas inmutables
    de ubicación / orden / tamaño de los logos del proyecto CC217."""


# ===========================================================================
# UTILIDADES XML
# ===========================================================================

def _set_cell_borders(cell, color="auto", size_pt=0):
    """Quita o pone bordes a una celda. size_pt=0 ⇒ sin borde."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil" if size_pt == 0 else "single")
        if size_pt > 0:
            b.set(qn("w:sz"), str(int(size_pt * 8)))
            b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _remove_table_borders(table):
    """Quita TODOS los bordes de una tabla (visualmente invisible)."""
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


def _force_fixed_layout(table):
    """Forzar layout fijo para que respete los anchos por celda exactamente."""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _add_horizontal_line(paragraph, color_hex="E94513", size_pt=1.0):
    """Añade una línea horizontal (border-bottom) bajo un párrafo."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_cell_margins(cell, top=20, bottom=20, left=30, right=30):
    """Márgenes internos de celda en twips (1 cm ≈ 567)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tc_mar.append(m)
    tc_pr.append(tc_mar)


# ===========================================================================
# CONFIGURACIÓN DE PÁGINA Y ESTILOS
# ===========================================================================

def _setup_page(doc):
    """A4 (21 × 29.7 cm), márgenes y header/footer estandarizados."""
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.top_margin    = Cm(3.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)


def _configure_default_styles(doc):
    """Tipografía FUCAI (Space Grotesk + Calibri)."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = NEGRO
    rpr = normal.element.get_or_add_rPr()
    existing_fonts = rpr.find(qn("w:rFonts"))
    r_fonts = existing_fonts if existing_fonts is not None else OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_BODY)
    r_fonts.set(qn("w:hAnsi"), FONT_BODY)
    r_fonts.set(qn("w:cs"), FONT_BODY)
    if existing_fonts is None:
        rpr.append(r_fonts)

    heading_specs = [
        ("Heading 1", 22, NARANJA_FUCAI, True),
        ("Heading 2", 16, NARANJA_FUCAI, True),
        ("Heading 3", 13, NEGRO,         True),
    ]
    for name, size, color, bold in heading_specs:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT_HEADING
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        rpr = st.element.get_or_add_rPr()
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT_HEADING)
        r_fonts.set(qn("w:hAnsi"), FONT_HEADING)
        rpr.append(r_fonts)


# ===========================================================================
# HEADER — exactamente 2 logos cuadrados idénticos
# ===========================================================================

def _build_header(doc, assets_dir):
    """
    Header de cada página:
      [AICS 2×2 cm | NAANE 2×2 cm]
      ── línea naranja FUCAI ──

    Las 2 celdas son IDÉNTICAS en tamaño. Los 2 logos son IDÉNTICOS en tamaño.
    """
    header = doc.sections[0].header
    for p in header.paragraphs:
        p.clear()

    # Tabla 1×2 invisible. Página útil 16 cm → 8 cm por columna.
    table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)
    _force_fixed_layout(table)

    column_width_cm = 8.0
    table.columns[0].width = Cm(column_width_cm)
    table.columns[1].width = Cm(column_width_cm)

    for i, filename in enumerate(HEADER_LOGOS_ORDER):
        cell = table.rows[0].cells[i]
        cell.width = Cm(column_width_cm)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)
        _set_cell_margins(cell, top=20, bottom=20, left=40, right=40)

        p = cell.paragraphs[0]
        # AICS (idx 0) alineado a la izquierda, NAANE (idx 1) alineado a la derecha
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        logo_path = Path(assets_dir) / "logos" / filename
        if not logo_path.exists():
            raise FileNotFoundError(f"Logo del header no encontrado: {logo_path}")

        run = p.add_run()
        # IMPORTANTE: width Y height idénticos para forzar cuadrado uniforme.
        # Los logos preprocesados son 1024×1024 px, por lo que un cuadrado
        # 2.0×2.0 cm les da tamaño visual idéntico sin deformación.
        run.add_picture(str(logo_path), width=HEADER_LOGO_SIDE, height=HEADER_LOGO_SIDE)

    # Línea naranja separadora debajo del header
    separator = header.add_paragraph()
    separator.paragraph_format.space_before = Pt(2)
    separator.paragraph_format.space_after = Pt(0)
    _add_horizontal_line(separator, color_hex="E94513", size_pt=1.0)


# ===========================================================================
# FOOTER — exactamente 9 logos cuadrados idénticos
# ===========================================================================

def _add_page_number_field(paragraph):
    """Inserta el campo dinámico 'Página X de Y'."""
    run = paragraph.add_run("Página ")
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS_MEDIO

    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText"); instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    r1 = paragraph.add_run()._r
    r1.append(fld_char1); r1.append(instr_text); r1.append(fld_char2)

    run2 = paragraph.add_run(" de ")
    run2.font.name = FONT_BODY
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRIS_MEDIO

    fld_char3 = OxmlElement("w:fldChar"); fld_char3.set(qn("w:fldCharType"), "begin")
    instr_text2 = OxmlElement("w:instrText"); instr_text2.text = "NUMPAGES"
    fld_char4 = OxmlElement("w:fldChar"); fld_char4.set(qn("w:fldCharType"), "end")
    r2 = paragraph.add_run()._r
    r2.append(fld_char3); r2.append(instr_text2); r2.append(fld_char4)


def _build_footer(doc, assets_dir):
    """
    Footer de cada página:
      ── línea naranja FUCAI ──
      [OIKOS][UniTrento][FUCAI][Javeriana][Alc.Leticia][ACITAM][ATICOYA][Maikuchiga][Raíz]
      Caption fina con identificadores AICS/OIKOS/FUCAI
      Página X de Y (derecha)

    Las 9 celdas tienen ANCHO IDÉNTICO. Los 9 logos tienen TAMAÑO IDÉNTICO (1.2×1.2 cm).
    """
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p.clear()

    # 1. Línea naranja superior
    line_para = footer.paragraphs[0]
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(2)
    _add_horizontal_line(line_para, color_hex="E94513", size_pt=0.75)

    # 2. Tabla 1×9 invisible con celdas de ANCHO IDÉNTICO
    n_logos = len(FOOTER_LOGOS_ORDER)  # = 9
    page_useful_width_cm = 16.0
    cell_width_cm = page_useful_width_cm / n_logos  # = 1.778 cm cada una

    logo_table = footer.add_table(rows=1, cols=n_logos, width=Cm(page_useful_width_cm))
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(logo_table)
    _force_fixed_layout(logo_table)

    for i, filename in enumerate(FOOTER_LOGOS_ORDER):
        cell = logo_table.rows[0].cells[i]
        cell.width = Cm(cell_width_cm)
        logo_table.columns[i].width = Cm(cell_width_cm)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_borders(cell)
        _set_cell_margins(cell, top=20, bottom=20, left=15, right=15)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        logo_path = Path(assets_dir) / "logos" / filename
        if not logo_path.exists():
            raise FileNotFoundError(f"Logo del footer no encontrado: {logo_path}")

        run = p.add_run()
        # IMPORTANTE: forzar tamaño CUADRADO idéntico para los 9 logos.
        # Los assets están preprocesados a 1024×1024 px con contenido centrado,
        # así que un cuadrado 1.2×1.2 cm les da a todos exactamente el mismo
        # tamaño visual sin deformar el contenido.
        run.add_picture(str(logo_path), width=FOOTER_LOGO_SIDE, height=FOOTER_LOGO_SIDE)

    # 3. Caption fina
    caption = footer.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(0)
    run_caption = caption.add_run(FOOTER_CAPTION)
    run_caption.font.name = FONT_BODY
    run_caption.font.size = Pt(7)
    run_caption.font.color.rgb = GRIS_MEDIO
    run_caption.italic = True

    # 4. Paginación
    page_para = footer.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_para.paragraph_format.space_before = Pt(0)
    page_para.paragraph_format.space_after = Pt(0)
    _add_page_number_field(page_para)


# ===========================================================================
# BLOQUE DE TÍTULO
# ===========================================================================

def _add_title_block(doc, title, subtitle, document_type, addressee):
    """Bloque inicial del cuerpo. Para cartas: fecha + destinatario.
    Para informes/actas/etc.: identificador CC217 + título + subtítulo."""
    if document_type in ("carta", "memorando"):
        from datetime import date
        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio",
                 "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
        hoy = date.today()
        fecha_str = f"Bogotá, D.C., {hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p_date.add_run(fecha_str)
        r.font.name = FONT_BODY
        r.font.size = Pt(11)

        doc.add_paragraph()

        if addressee:
            for line in [
                addressee.get("name", ""),
                addressee.get("position", ""),
                addressee.get("organization", ""),
                addressee.get("city", ""),
            ]:
                if line:
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.name = FONT_BODY
                    r.font.size = Pt(11)
                    if line == addressee.get("name", ""):
                        r.bold = True
            doc.add_paragraph()

        # Mostrar identificador CC217 en el asunto para que el validador lo encuentre
        if title:
            p_asunto = doc.add_paragraph()
            r1 = p_asunto.add_run("Asunto: ")
            r1.bold = True; r1.font.name = FONT_BODY; r1.font.size = Pt(11)
            r2 = p_asunto.add_run(f"{title} — Proyecto {PROJECT_IDENTIFIER}")
            r2.font.name = FONT_BODY; r2.font.size = Pt(11)
            doc.add_paragraph()
    else:
        p_id = doc.add_paragraph()
        p_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_id = p_id.add_run(PROJECT_NAME_FULL)
        r_id.font.name = FONT_BODY
        r_id.font.size = Pt(10)
        r_id.font.color.rgb = GRIS_MEDIO
        r_id.italic = True
        p_id.paragraph_format.space_after = Pt(2)

        p_title = doc.add_paragraph()
        r_title = p_title.add_run(title)
        r_title.font.name = FONT_HEADING
        r_title.font.size = Pt(22)
        r_title.font.color.rgb = NARANJA_FUCAI
        r_title.bold = True
        p_title.paragraph_format.space_after = Pt(4)

        if subtitle:
            p_sub = doc.add_paragraph()
            r_sub = p_sub.add_run(subtitle)
            r_sub.font.name = FONT_HEADING
            r_sub.font.size = Pt(14)
            r_sub.font.color.rgb = NEGRO
            p_sub.paragraph_format.space_after = Pt(12)

        sep = doc.add_paragraph()
        _add_horizontal_line(sep, color_hex="E94513", size_pt=0.5)
        sep.paragraph_format.space_after = Pt(12)


# ===========================================================================
# VALIDACIÓN AUTOMÁTICA POST-GENERACIÓN
# ===========================================================================

def validate_naane_document(docx_path, assets_dir=None):
    """
    Verifica que el documento .docx cumpla las reglas inmutables del membrete
    NAANE / CC217:

      1. El header tiene EXACTAMENTE 2 imágenes embebidas.
      2. El footer tiene EXACTAMENTE 9 imágenes embebidas.
      3. Las imágenes del header y footer aparecen en el ORDEN CANÓNICO esperado.
         (Verificado por hash MD5 del contenido binario, no por nombre, porque
         python-docx renombra los archivos a image1.png, image2.png, etc.)
      4. Las 2 imágenes del header tienen DIMENSIONES IDÉNTICAS.
      5. Las 9 imágenes del footer tienen DIMENSIONES IDÉNTICAS.
      6. El cuerpo contiene la cadena 'CC217' en al menos un párrafo.

    Parámetros
    ----------
    docx_path : str | Path
        Ruta al documento .docx generado.
    assets_dir : str | Path | None
        Ruta al directorio assets/ del skill, para poder calcular los hashes
        esperados de los logos canónicos. Si es None, se intenta inferir
        desde la ubicación de este script.

    Lanza NaaneBrandingValidationError si algo falla.
    Retorna un dict con el resumen de la verificación si todo está bien.
    """
    import zipfile
    import re
    import hashlib
    from collections import OrderedDict

    docx_path = str(docx_path)

    # Resolver assets_dir
    if assets_dir is None:
        assets_dir = Path(__file__).resolve().parent.parent / "assets"
    assets_dir = Path(assets_dir)

    # Pre-calcular hashes MD5 de los logos canónicos
    def _md5(path):
        h = hashlib.md5()
        with open(path, 'rb') as f:
            h.update(f.read())
        return h.hexdigest()

    expected_header_hashes = [
        _md5(assets_dir / "logos" / fname) for fname in HEADER_LOGOS_ORDER
    ]
    expected_footer_hashes = [
        _md5(assets_dir / "logos" / fname) for fname in FOOTER_LOGOS_ORDER
    ]
    hash_to_name = {}
    for fname in HEADER_LOGOS_ORDER + FOOTER_LOGOS_ORDER:
        hash_to_name[_md5(assets_dir / "logos" / fname)] = fname

    report = OrderedDict()
    errors = []

    with zipfile.ZipFile(docx_path, 'r') as zf:
        names = zf.namelist()

        # Construir mapa de imagen embebida → hash MD5
        embedded_images = {}  # 'media/imageN.png' → md5 hash
        for n in names:
            if n.startswith('word/media/'):
                content = zf.read(n)
                embedded_images[Path(n).name] = hashlib.md5(content).hexdigest()

        # --- 1. Header ---
        header_files = [n for n in names if re.match(r'word/header\d+\.xml$', n)]
        if not header_files:
            errors.append("No se encontró ningún archivo header*.xml en el documento.")
        else:
            header_xml = zf.read(header_files[0]).decode('utf-8')
            header_rels_path = f"word/_rels/{Path(header_files[0]).name}.rels"
            header_rels = zf.read(header_rels_path).decode('utf-8') if header_rels_path in names else ""

            drawings_in_header = len(re.findall(r'<w:drawing>', header_xml))
            report["header_drawings_count"] = drawings_in_header
            if drawings_in_header != len(HEADER_LOGOS_ORDER):
                errors.append(
                    f"Header debe tener exactamente {len(HEADER_LOGOS_ORDER)} imágenes "
                    f"(AICS + NAANE), pero tiene {drawings_in_header}."
                )

            header_extents = re.findall(
                r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', header_xml
            )
            report["header_image_dimensions_emu"] = header_extents
            if len(set(header_extents)) > 1:
                errors.append(
                    f"Las imágenes del header NO tienen dimensiones idénticas: {header_extents}"
                )

            # Verificar orden por HASH de contenido
            header_embed_ids = re.findall(r'r:embed="(rId\d+)"', header_xml)
            header_rels_map = dict(re.findall(
                r'Id="(rId\d+)"\s+Type="[^"]*image"\s+Target="([^"]+)"',
                header_rels
            ))
            header_actual_hashes = []
            header_resolved_names = []
            for rid in header_embed_ids:
                target_name = Path(header_rels_map.get(rid, "")).name
                actual_hash = embedded_images.get(target_name, "")
                header_actual_hashes.append(actual_hash)
                # Resolver nombre canónico via hash
                canonical_name = hash_to_name.get(actual_hash, f"DESCONOCIDO({target_name})")
                header_resolved_names.append(canonical_name)

            report["header_order_resolved"] = header_resolved_names
            report["header_order_expected"] = list(HEADER_LOGOS_ORDER)
            if header_actual_hashes != expected_header_hashes:
                errors.append(
                    "El ORDEN de los logos del header no coincide con el canónico.\n"
                    f"  Esperado: {list(HEADER_LOGOS_ORDER)}\n"
                    f"  Encontrado: {header_resolved_names}"
                )

        # --- 2. Footer ---
        footer_files = [n for n in names if re.match(r'word/footer\d+\.xml$', n)]
        if not footer_files:
            errors.append("No se encontró ningún archivo footer*.xml en el documento.")
        else:
            footer_xml = zf.read(footer_files[0]).decode('utf-8')
            footer_rels_path = f"word/_rels/{Path(footer_files[0]).name}.rels"
            footer_rels = zf.read(footer_rels_path).decode('utf-8') if footer_rels_path in names else ""

            drawings_in_footer = len(re.findall(r'<w:drawing>', footer_xml))
            report["footer_drawings_count"] = drawings_in_footer
            if drawings_in_footer != len(FOOTER_LOGOS_ORDER):
                errors.append(
                    f"Footer debe tener exactamente {len(FOOTER_LOGOS_ORDER)} imágenes, "
                    f"pero tiene {drawings_in_footer}."
                )

            footer_extents = re.findall(
                r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', footer_xml
            )
            report["footer_image_dimensions_emu"] = footer_extents
            if len(set(footer_extents)) > 1:
                errors.append(
                    f"Las imágenes del footer NO tienen dimensiones idénticas. "
                    f"Dimensiones encontradas: {set(footer_extents)}"
                )

            # Verificar orden por HASH
            footer_embed_ids = re.findall(r'r:embed="(rId\d+)"', footer_xml)
            footer_rels_map = dict(re.findall(
                r'Id="(rId\d+)"\s+Type="[^"]*image"\s+Target="([^"]+)"',
                footer_rels
            ))
            footer_actual_hashes = []
            footer_resolved_names = []
            for rid in footer_embed_ids:
                target_name = Path(footer_rels_map.get(rid, "")).name
                actual_hash = embedded_images.get(target_name, "")
                footer_actual_hashes.append(actual_hash)
                canonical_name = hash_to_name.get(actual_hash, f"DESCONOCIDO({target_name})")
                footer_resolved_names.append(canonical_name)

            report["footer_order_resolved"] = footer_resolved_names
            report["footer_order_expected"] = list(FOOTER_LOGOS_ORDER)
            if footer_actual_hashes != expected_footer_hashes:
                errors.append(
                    "El ORDEN de los logos del footer no coincide con el canónico.\n"
                    f"  Esperado: {list(FOOTER_LOGOS_ORDER)}\n"
                    f"  Encontrado: {footer_resolved_names}"
                )

        # --- 3. Body: identificador CC217 ---
        doc_xml = zf.read('word/document.xml').decode('utf-8')
        body_text = " ".join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', doc_xml))
        if PROJECT_IDENTIFIER not in body_text:
            errors.append(
                f"El cuerpo del documento NO contiene el identificador "
                f"obligatorio '{PROJECT_IDENTIFIER}'."
            )
        report["body_contains_cc217"] = (PROJECT_IDENTIFIER in body_text)

    if errors:
        msg = "Validación del membrete NAANE / CC217 FALLÓ:\n" + "\n".join(
            f"  ✗ {e}" for e in errors
        )
        msg += f"\n\nReporte completo:\n{dict(report)}"
        raise NaaneBrandingValidationError(msg)

    report["status"] = "OK"
    return report


# ===========================================================================
# API PÚBLICA
# ===========================================================================

def create_naane_document(
    title: str,
    subtitle: str = "",
    document_type: str = "informe",
    addressee: dict = None,
    skill_dir: str = None,
):
    """
    Crea un Document() python-docx con el membrete oficial CC217 — Naãne.

    Parámetros
    ----------
    title : str
        Título principal del documento (o asunto si es carta).
    subtitle : str
        Subtítulo opcional.
    document_type : str
        Uno de: 'informe', 'carta', 'acta', 'presupuesto', 'memorando', 'concepto'.
    addressee : dict | None
        Solo para cartas. Claves: name, position, organization, city.
    skill_dir : str
        Ruta absoluta al directorio del skill (donde está assets/logos/).

    Retorna
    -------
    docx.Document — documento listo para agregar contenido y guardar.

    Para validar el documento después de guardarlo:
        doc.save("/path/to/file.docx")
        validate_naane_document("/path/to/file.docx")
    """
    if skill_dir is None:
        skill_dir = str(Path(__file__).resolve().parent.parent)

    assets_dir = Path(skill_dir) / "assets"
    if not assets_dir.exists():
        raise FileNotFoundError(
            f"No se encuentra assets/ en {skill_dir}. Verifica la ruta del skill."
        )

    # Verificar que TODOS los logos requeridos existen ANTES de empezar
    missing = []
    for fname in HEADER_LOGOS_ORDER + FOOTER_LOGOS_ORDER:
        if not (assets_dir / "logos" / fname).exists():
            missing.append(fname)
    if missing:
        raise FileNotFoundError(
            f"Logos faltantes en {assets_dir / 'logos'}: {missing}"
        )

    doc = Document()
    _setup_page(doc)
    _configure_default_styles(doc)
    _build_header(doc, assets_dir)
    _build_footer(doc, assets_dir)
    _add_title_block(doc, title, subtitle, document_type, addressee)

    return doc


# ===========================================================================
# CLI / prueba local
# ===========================================================================

if __name__ == "__main__":
    import sys
    skill_dir = sys.argv[1] if len(sys.argv) > 1 else str(Path(__file__).resolve().parent.parent)

    doc = create_naane_document(
        title="Informe trimestral de avance",
        subtitle="Periodo enero–marzo 2026",
        document_type="informe",
        skill_dir=skill_dir,
    )

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Durante el primer trimestre del año 2026, el proyecto «Naãne — Buen Vivir "
        "Amazonas» avanzó en la consolidación de los procesos de gobernanza "
        "territorial con las comunidades indígenas del Trapecio Amazónico."
    )

    doc.add_heading("2. Avances por componente", level=1)
    doc.add_heading("2.1. Gobernanza territorial", level=2)
    doc.add_paragraph(
        "Se acompañaron seis cabildos en la formulación de sus planes propios "
        "de vida, articulando con ACITAM y ATICOYA los mecanismos de coordinación."
    )

    doc.add_heading("2.2. Seguridad alimentaria y soberanía", level=2)
    doc.add_paragraph(
        "Los huertos comunitarios y chagras tradicionales en doce comunidades "
        "reportan un incremento del 32 % en la diversidad de especies cultivadas."
    )

    doc.add_heading("3. Articulación interinstitucional", level=1)
    doc.add_paragraph(
        "El consorcio avanzó en reuniones con OIKOS (coordinación técnica desde Italia), "
        "la Pontificia Universidad Javeriana, el Dipartimento di Ingegneria de UniTrento, "
        "y las fundaciones Maikuchiga y Raíz."
    )

    out = Path(skill_dir) / "test_output.docx"
    doc.save(str(out))
    print(f"Documento generado: {out}")

    print("\n--- Ejecutando validación post-generación ---")
    try:
        report = validate_naane_document(out)
        print("✓ Validación EXITOSA")
        for k, v in report.items():
            print(f"  {k}: {v}")
    except NaaneBrandingValidationError as e:
        print(f"✗ Validación FALLÓ:\n{e}")
        sys.exit(1)
